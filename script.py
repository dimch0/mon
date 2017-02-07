#!/usr/bin/env python2
# -*- coding: utf-8 -*-
import sys
from docx import Document
from docx.shared import Pt
from xlrd import open_workbook

FOR_AGREEMENT_FILE = 'for-agreements.xlsx'
DOG_TEMPLATE = "contract_template.docx"
DOG_NEW = u"Договор-за-обучение_BG-EN_Appendix-B_HRC-Foundation-002_{0}.docx"

IDX_NAME_BG = 0
IDX_NATIONAL_ID = 1
IDX_ID_CARD_NO = 2
IDX_ISSUED_ON = 3
IDX_ISSUED_BY = 4
IDX_ADDRESS_BG = 5
IDX_PHONE = 6
IDX_NAME_E = 7
IDX_ISSUED_BY_E =8
IDX_ADDRESS_E = 9


def write_doc(doc_template, new_doc, info_list):

    doc = Document(doc_template)
    table = doc.tables[0]

    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(11)
    font.name = 'Cambria'

    name_bg_here = info_list[0]
    name_e_here = info_list[1]
    info_bg = info_list[2]
    info_e = info_list[3]
    adr_bg = info_list[4]
    adr_e = info_list[5]

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:

                if u"NAME_BG_HERE" in paragraph.text:
                    paragraph.text = paragraph.text.replace("NAME_BG_HERE", "")
                    paragraph.add_run(name_bg_here).bold = True
                    paragraph.style = doc.styles['Normal']
                if u"NAME_E_HERE" in paragraph.text:
                    paragraph.text = paragraph.text.replace("NAME_E_HERE", "")
                    paragraph.add_run(name_e_here).bold = True
                    paragraph.style = doc.styles['Normal']
                if u"INFO_BG" in paragraph.text:
                    paragraph.text = paragraph.text.replace("INFO_BG", info_bg)
                    paragraph.style = doc.styles['Normal']
                if u"INFO_E" in paragraph.text:
                    paragraph.text = paragraph.text.replace("INFO_E", info_e)
                    paragraph.style = doc.styles['Normal']
                if u"ADDRESS_BG" in paragraph.text:
                    paragraph.text = paragraph.text.replace("ADDRESS_BG", adr_bg)
                    paragraph.style = doc.styles['Normal']
                if u"ADDRRES_E" in paragraph.text:
                    paragraph.text = paragraph.text.replace("ADDRRES_E", adr_e)
                    paragraph.style = doc.styles['Normal']

    doc.save(new_doc)


def main():
    wb = open_workbook(FOR_AGREEMENT_FILE)

    for sheet in wb.sheets():
        number_of_rows = sheet.nrows

        for row in range(1, number_of_rows):

            NAME_BG = sheet.cell(row,IDX_NAME_BG).value if sheet.cell(row,IDX_NAME_BG).value else ""
            NATIONAL_ID = str(int(sheet.cell(row,IDX_NATIONAL_ID).value)) if sheet.cell(row,IDX_NATIONAL_ID).value else ""
            ID_CARD_NO = str(int(sheet.cell(row,IDX_ID_CARD_NO).value)) if sheet.cell(row,IDX_ID_CARD_NO).value else ""
            ISSUED_ON = str(sheet.cell(row,IDX_ISSUED_ON).value).replace("/", ".").replace(",", ".") if sheet.cell(row,IDX_ISSUED_ON).value else ""
            ISSUED_BY = sheet.cell(row,IDX_ISSUED_BY).value if sheet.cell(row,IDX_ISSUED_BY).value else ""
            ADDRESS_BG = sheet.cell(row,IDX_ADDRESS_BG).value if sheet.cell(row,IDX_ADDRESS_BG).value else ""
            ADDRESS_E = sheet.cell(row,IDX_ADDRESS_E).value if sheet.cell(row,IDX_ADDRESS_E).value else ""
            PHONE = str(sheet.cell(row,IDX_PHONE).value) if sheet.cell(row,IDX_PHONE).value else ""
            NAME_E = sheet.cell(row,IDX_NAME_E).value if sheet.cell(row,IDX_NAME_E).value else ""
            ISSUED_BY_E = sheet.cell(row,IDX_ISSUED_BY_E).value if sheet.cell(row,IDX_ISSUED_BY_E).value else ""

            INFO_BG = u"{national_id}{ID_CARD_NO}{ISSUED_ON}{ISSUED_BY}".format(
                national_id = u"с ЕГН {0}, ".format(NATIONAL_ID) if NATIONAL_ID else "",
                ID_CARD_NO = u"л.к. {0}, ".format(ID_CARD_NO) if ID_CARD_NO else "",
                ISSUED_ON = u"изд. на {0}".format(ISSUED_ON) if ISSUED_ON else "",
                ISSUED_BY = u" от МВР – {0}".format(ISSUED_BY) if ISSUED_BY else "")

            INFO_E = u"{national_id}{ID_CARD_NO}{ISSUED_ON}{ISSUED_BY}".format(
                national_id = u"with National ID No. {0}, ".format(NATIONAL_ID) if NATIONAL_ID else "",
                ID_CARD_NO = u"personal ID card No {0}, ".format(ID_CARD_NO) if ID_CARD_NO else "",
                ISSUED_ON = u"issued on {0}".format(ISSUED_ON) if ISSUED_ON else "",
                ISSUED_BY = u" by the Ministry of Interior – {0}".format(ISSUED_BY_E) if ISSUED_BY_E else "")

            ADDRESS_BG = u"Адрес по регистрация: {ADDRESS}{PHONE}".format(
                ADDRESS = ADDRESS_BG,
                PHONE = u", тел. {0}".format(PHONE) if PHONE else "")
            ADDRRESS_E = u"Registered address: {ADDRESS}{PHONE}".format(
                ADDRESS = ADDRESS_E,
                PHONE = u", tel. {0}".format(PHONE) if PHONE else "")

            info_list = [NAME_BG, NAME_E, INFO_BG, INFO_E, ADDRESS_BG, ADDRRESS_E]

            new_doc_name = DOG_NEW.format("_".join(NAME_E.split()))
            print u"Creating document: {0} ...".format(new_doc_name)
            write_doc(DOG_TEMPLATE, new_doc_name, info_list)


if __name__ == '__main__':
    main()
    sys.exit()
